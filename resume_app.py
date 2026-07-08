"""
简历生成器 Web 应用 — Streamlit Cloud 版
支持上传 .doc/.docx，提取数据 → 编辑 → 生成 PDF
V4 · IBM×Apple×Stripe 设计语言 · 央企项目经理竞聘专用
"""
import streamlit as st
import json, os, re, base64, tempfile, subprocess
from pathlib import Path

st.set_page_config(page_title="央企项目经理竞聘简历生成器", page_icon="📄", layout="wide")

# ── CSS ──
BASE_CSS = """
:root {
  --navy: #1A2332; --navy-light: #2D3748; --red: #C41230; --blue: #2B5C9E;
  --blue-light: #EBF1F8; --gold: #B8860B; --gold-light: #FBF5E8;
  --white: #FFFFFF; --surface: #F7F8FA; --border: #E2E8F0; --border-light: #EDF2F7;
  --text-primary: #1A202C; --text-body: #4A5568; --text-muted: #718096;
  --shadow-sm: 0 1px 3px rgba(26,35,50,0.06);
  --font: "Microsoft YaHei", "微软雅黑", "PingFang SC", "Noto Sans SC", sans-serif;
  --font-light: "Microsoft YaHei Light", "PingFang SC Light", sans-serif;
}
@page { size: A4; margin: 0; }
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: var(--font); font-size: 9pt; line-height: 1.6; color: var(--text-body); background: var(--white); -webkit-print-color-adjust: exact; print-color-adjust: exact; }
.page-layout { display: flex; width: 210mm; min-height: 297mm; page-break-after: always; }
.page-layout:last-of-type { page-break-after: auto; }
.left-panel { width: 32%; background: linear-gradient(175deg, #1A2332 0%, #1E2D3D 100%); color: rgba(255,255,255,0.85); padding: 32px 20px 28px 22px; position: relative; overflow: hidden; }
.left-panel::before { content: ""; position: absolute; top: -60px; right: -60px; width: 200px; height: 200px; background: radial-gradient(circle, rgba(255,255,255,0.025) 0%, transparent 70%); pointer-events: none; }
.lp-photo { width: 72px; height: 96px; border: 1.5px solid rgba(255,255,255,0.2); border-radius: 4px; margin-bottom: 18px; overflow: hidden; background: rgba(255,255,255,0.04); }
.lp-photo img { width: 100%; height: 100%; object-fit: cover; display: block; }
.lp-name { font-size: 24px; font-weight: 300; letter-spacing: 3px; font-family: var(--font-light); color: var(--white); }
.lp-title { font-size: 9px; font-weight: 400; color: rgba(255,255,255,0.65); margin-top: 4px; letter-spacing: 1px; line-height: 1.5; }
.lp-divider { width: 28px; height: 2px; background: var(--gold); margin: 14px 0; }
.lp-section { margin-top: 18px; }
.lp-section h3 { font-size: 9px; font-weight: 700; letter-spacing: 1.5px; color: var(--gold); margin-bottom: 7px; text-transform: uppercase; }
.lp-section .lp-item { font-size: 7.5px; color: rgba(255,255,255,0.7); line-height: 1.9; }
.lp-contact { margin-top: 22px; padding-top: 14px; border-top: 1px solid rgba(255,255,255,0.1); }
.lp-contact .lp-contact-item { font-size: 7px; color: rgba(255,255,255,0.5); line-height: 2; }
.left-panel .skill-row .skill-name { color: rgba(255,255,255,0.8); }
.left-panel .skill-row .skill-bar { background: rgba(255,255,255,0.15); }
.right-panel { flex: 1; padding: 28px 22px 24px 24px; }
.right-section { margin-bottom: 16px; }
.right-section h2 { font-size: 11px; font-weight: 700; color: var(--navy); padding-bottom: 6px; border-bottom: 2px solid var(--border); margin-bottom: 10px; letter-spacing: 1px; display: flex; align-items: center; gap: 8px; }
.right-section h2 .sec-icon { display: inline-flex; align-items: center; justify-content: center; min-width: 18px; height: 18px; background: var(--navy); color: var(--white); font-size: 8px; font-weight: 700; border-radius: 3px; }
.one-liner-box { padding: 10px 14px; background: var(--blue-light); border-left: 3px solid var(--blue); border-radius: 0 4px 4px 0; margin-bottom: 16px; }
.one-liner-box p { font-size: 8px; color: var(--navy-light); line-height: 1.6; }
.tl-item { position: relative; padding-left: 15px; margin-bottom: 10px; border-left: 1.5px solid var(--border); }
.tl-item::before { content: ""; position: absolute; left: -4px; top: 1px; width: 6.5px; height: 6.5px; border-radius: 50%; background: var(--blue); }
.tl-item .tl-date { font-size: 7.5px; font-weight: 700; color: var(--blue); letter-spacing: 0.3px; }
.tl-item .tl-head { font-size: 9.5px; font-weight: 700; color: var(--text-primary); }
.tl-item .tl-sub { font-size: 7.5px; color: var(--text-muted); line-height: 1.5; }
.tag { display: inline-block; font-size: 6px; font-weight: 700; padding: 2px 7px; border-radius: 10px; letter-spacing: 0.5px; }
.tag-blue { background: var(--blue-light); color: var(--blue); }
.tag-orange { background: var(--gold-light); color: var(--gold); }
.tag-green { background: #E8F5E9; color: #2E7D32; }
.tag-red { background: #FDEAEC; color: var(--red); }
.proj-card { background: var(--surface); border: 1px solid var(--border-light); border-radius: 5px; padding: 10px 12px; margin-bottom: 8px; box-shadow: var(--shadow-sm); }
.proj-card .proj-top { display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 3px; }
.proj-card .proj-name { font-size: 10px; font-weight: 700; color: var(--text-primary); }
.proj-card .proj-meta { font-size: 7px; color: var(--text-muted); margin-bottom: 5px; }
.proj-card .proj-metrics { display: flex; gap: 14px; padding: 5px 0; margin-bottom: 5px; border-top: 1px solid var(--border-light); border-bottom: 1px solid var(--border-light); }
.proj-card .proj-metric-val { font-size: 12px; font-weight: 700; color: var(--navy); }
.proj-card .proj-metric-label { font-size: 6px; color: var(--text-muted); display: block; }
.proj-card .proj-dots { list-style: none; font-size: 7.5px; color: var(--text-body); line-height: 1.6; }
.proj-card .proj-dots li { padding: 1px 0 1px 9px; position: relative; }
.proj-card .proj-dots li::before { content: ""; position: absolute; left: 0; top: 4px; width: 3.5px; height: 3.5px; border-radius: 50%; background: var(--blue); }
.dashboard { display: flex; gap: 8px; margin-bottom: 16px; }
.dash-kpi { flex: 1; text-align: center; padding: 10px 6px; border-radius: 5px; color: var(--white); box-shadow: var(--shadow-sm); }
.dash-kpi .kpi-val { font-size: 20px; font-weight: 700; line-height: 1.1; letter-spacing: -0.5px; }
.dash-kpi .kpi-label { font-size: 6.5px; margin-top: 3px; opacity: 0.85; letter-spacing: 0.5px; }
.dash-navy { background: var(--navy); } .dash-red { background: var(--red); } .dash-blue { background: var(--blue); } .dash-gold { background: #8B6914; }
.dual-track { display: flex; gap: 8px; }
.track-box { flex: 1; padding: 10px; border-radius: 4px; border-top: 3px solid; }
.track-box h3 { font-size: 9px; font-weight: 700; margin-bottom: 6px; }
.track-box ul { list-style: none; font-size: 7.5px; color: var(--text-body); line-height: 1.6; }
.track-box ul li { padding: 1.5px 0 1.5px 8px; position: relative; }
.track-box ul li::before { content: ""; position: absolute; left: 0; top: 4px; width: 3px; height: 3px; border-radius: 50%; }
.track-business { background: var(--gold-light); border-color: var(--gold); }
.track-business h3 { color: #8B6914; } .track-business li::before { background: var(--gold); }
.track-delivery { background: var(--blue-light); border-color: var(--blue); }
.track-delivery h3 { color: var(--blue); } .track-delivery li::before { background: var(--blue); }
.ai-section { background: linear-gradient(135deg, #F0F4F8 0%, #F7F8FA 100%); border: 1px solid var(--border-light); border-radius: 5px; padding: 10px 14px; }
.ai-section h3 { font-size: 9px; font-weight: 700; color: var(--navy); margin-bottom: 3px; }
.skill-row { display: flex; align-items: center; margin-bottom: 5px; gap: 8px; }
.skill-row .skill-name { min-width: 56px; font-size: 7.5px; font-weight: 700; color: var(--text-primary); text-align: right; }
.skill-row .skill-bar { flex: 1; height: 4px; background: var(--border-light); border-radius: 2px; overflow: hidden; }
.skill-row .skill-fill { height: 100%; border-radius: 2px; }
.skill-fill.navy { background: var(--navy); } .skill-fill.blue { background: var(--blue); } .skill-fill.gold { background: var(--gold); } .skill-fill.red { background: var(--red); }
.page-break { page-break-before: always; }
"""

# ── 模板 ──
RESUME_TEMPLATE = """<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8"><title>{{ name }} — 竞聘简历</title><style>""" + BASE_CSS + """</style></head><body>
<div class="page-layout">
  <div class="left-panel">
    <div class="lp-photo">{% if photo %}<img src="{{ photo }}" alt="照片">{% endif %}</div>
    <div class="lp-name">{{ name }}</div><div class="lp-title">{{ title }}</div><div class="lp-divider"></div>
    <div class="lp-section"><h3>基本信息</h3><div class="lp-item">{% if gender %}性别：{{ gender }}<br>{% endif %}{% if birth %}出生：{{ birth }}<br>{% endif %}{% if political %}政治面貌：{{ political }}<br>{% endif %}{% if hometown %}籍贯：{{ hometown }}<br>{% endif %}{% if location %}工作地：{{ location }}{% endif %}</div></div>
    <div class="lp-section"><h3>教育背景</h3><div class="lp-item">{% for edu in education %}{{ edu.school }}<br><span style="color:rgba(255,255,255,0.45);font-size:7px;">{{ edu.major }} · {{ edu.degree }}<br>{{ edu.period }}</span>{% if not loop.last %}<br>{% endif %}{% endfor %}</div></div>
    {% if certifications %}<div class="lp-section"><h3>资质证书</h3><div class="lp-item">{% for cert in certifications %}{{ cert.name }}{% if cert.year %}（{{ cert.year }}）{% endif %}{% if not loop.last %}<br>{% endif %}{% endfor %}</div></div>{% endif %}
    <div class="lp-section"><h3>核心能力</h3>{% for sk in skills %}<div class="skill-row"><span class="skill-name">{{ sk.name }}</span><span class="skill-bar"><span class="skill-fill {{ sk.color }}" style="width:{{ sk.pct }}%;"></span></span></div>{% endfor %}</div>
    <div class="lp-contact"><div class="lp-contact-item">{% if phone %}📞 {{ phone }}<br>{% endif %}{% if email %}✉ {{ email }}<br>{% endif %}{% if location %}📍 {{ location }}{% endif %}</div></div>
  </div>
  <div class="right-panel">
    <div class="one-liner-box"><p>{{ one_liner }}</p></div>
    <div class="dashboard">
      <div class="dash-kpi dash-navy"><div class="kpi-val">{{ stats.total_contract }}</div><div class="kpi-label">管理合同额 / 亿元</div></div>
      <div class="dash-kpi dash-red"><div class="kpi-val">{{ stats.experience_years }}</div><div class="kpi-label">从业年限 / 年</div></div>
      <div class="dash-kpi dash-blue"><div class="kpi-val">{{ stats.project_count }}</div><div class="kpi-label">管理项目 / 个</div></div>
      <div class="dash-kpi dash-gold"><div class="kpi-val">{{ stats.team_size }}</div><div class="kpi-label">团队规模 / 人</div></div>
    </div>
    <div class="right-section"><h2><span class="sec-icon">01</span> 工作经历</h2>{% for exp in experiences %}<div class="tl-item"><div class="tl-date">{{ exp.period }}</div><div class="tl-head">{{ exp.role }}</div><div class="tl-sub">{{ exp.company }}{% if exp.summary %} — {{ exp.summary }}{% endif %}</div></div>{% endfor %}</div>
    <div class="right-section"><h2><span class="sec-icon">02</span> 代表性项目</h2>{% for proj in projects %}<div class="proj-card"><div class="proj-top"><div class="proj-name">{{ proj.name }}</div><span class="tag {% if '盾构' in proj.method %}tag-blue{% elif 'TBM' in proj.method %}tag-orange{% elif '矿山' in proj.method %}tag-green{% else %}tag-red{% endif %}">{{ proj.method }}</span></div><div class="proj-meta">{{ proj.role }} &nbsp;|&nbsp; {{ proj.period }}</div>{% if proj.metrics %}<div class="proj-metrics">{% for m in proj.metrics %}<div><span class="proj-metric-val">{{ m.value }}</span><span class="proj-metric-label">{{ m.label }}</span></div>{% endfor %}</div>{% endif %}<ul class="proj-dots">{% for h in proj.highlights %}<li>{{ h }}</li>{% endfor %}</ul></div>{% endfor %}</div>
  </div>
</div>
<div class="page-break"></div>
<div class="page-layout">
  <div class="left-panel">
    <div class="lp-section"><h3>岗位定位</h3><div class="lp-item">双核驱动：对内抓履约 + 对外拓市场<br><br>覆盖工法：盾构 · TBM · 矿山法 · 明挖法<br><br>团队统筹：七大职能部门 + 专业作业队 + 一线施工班组</div></div>
    <div class="lp-section"><h3>AI 数智化赋能</h3><div class="lp-item">{% for s in ai_scenarios %}{{ s.icon }}. {{ s.name }}<br><span style="color:rgba(255,255,255,0.4);font-size:6.5px;padding-left:10px;">{{ s.desc }}</span>{% if not loop.last %}<br>{% endif %}{% endfor %}</div></div>
  </div>
  <div class="right-panel">
    <div class="right-section"><h2><span class="sec-icon">03</span> 双核驱动能力</h2><div class="dual-track"><div class="track-box track-business"><h3>对外拓市场 — 经营线</h3><ul>{% for item in track_business %}<li>{{ item }}</li>{% endfor %}</ul></div><div class="track-box track-delivery"><h3>对内抓履约 — 管理线</h3><ul>{% for item in track_delivery %}<li>{{ item }}</li>{% endfor %}</ul></div></div></div>
    <div class="ai-section"><h3>个人发展理念</h3><p style="font-size:7.5px;color:var(--text-body);line-height:1.7;">{{ closing }}</p></div>
  </div>
</div>
</body></html>"""

# ── 默认数据 ──
DEFAULT_DATA = {
    "name": "", "title": "", "one_liner": "", "photo": "",
    "gender": "", "birth": "", "hometown": "", "political": "", "phone": "", "email": "", "location": "",
    "stats": {"experience_years": "", "total_contract": "", "project_count": "", "team_size": ""},
    "track_business": ["", "", ""],
    "track_delivery": ["", "", ""],
    "experiences": [],
    "projects": [],
    "education": [],
    "certifications": [],
    "awards": [], "patents": [],
    "skills": [
        {"name": "项目管理", "color": "navy", "pct": 95},
        {"name": "市场经营", "color": "gold", "pct": 85},
        {"name": "技术创新", "color": "blue", "pct": 80},
        {"name": "AI 数智化", "color": "red", "pct": 90}
    ],
    "ai_scenarios": [
        {"icon": "1", "name": "日报/周报自动生成", "desc": "3-5小时 → 15分钟"},
        {"icon": "2", "name": "会议纪要待办追踪", "desc": "1-2小时 → 15分钟"},
        {"icon": "3", "name": "招标情报每日监测", "desc": "2小时 → 12分钟"},
        {"icon": "4", "name": "施组方案智能起草", "desc": "编制周期缩短30%+"},
        {"icon": "5", "name": "工法专利报告起草", "desc": "周期缩短50%+"}
    ],
    "closing": "以履约立身，以经营致远，以创新破局。",
}

# 示例模板（何进江——基建央企项目经理竞聘数据，供参考/快速体验）
SAMPLE_DATA = {
    "name": "何进江",
    "title": "中铁隧道集团二处有限公司 · 党支部书记、项目经理",
    "one_liner": "兰州交通大学土木工程本科，17年隧道与轨道交通工程管理经验，历任南京地铁七号线、北京地铁1号线支线、深圳西丽站等重大项目项目经理，累计管理合同额超42亿元，具备盾构/矿山法/明挖法全工法管理能力，持有一级建造师（市政公用工程），清华大学卓越领袖研修班（2026年9月结业）",
    "photo": "",
    "gender": "男", "birth": "1985年2月", "hometown": "", "political": "", "phone": "", "email": "", "location": "福建厦泉/深圳",
    "stats": {"experience_years": "17", "total_contract": "42+", "project_count": "3", "team_size": "100+/2000+"},
    "track_business": [
        "福建片区区域经营负责人，统筹市场开发与客户关系维护",
        "深圳西丽站合同额32.76亿元，二处分劈19.72亿元，A类重点项目",
        "南京地铁清凉山站优化围护结构方案，节约成本2300万元",
        "三个项目责任成本盈利达标率100%，累计超责任成本盈利568.81万元"
    ],
    "track_delivery": [
        "累计管理合同额超42亿元，涵盖地铁车站（明挖/暗挖）、盾构隧道、矿山法隧道全类型",
        "南京清凉山站独拱两层岛式车站，大跨断面303㎡，优化工序节约工期约310天",
        "北京体育场南街站双层三跨矩形框架结构，地下连续墙+钢支撑明挖法",
        "深圳西丽站六大施工难点攻坚：盾构下穿建构筑物、小净距隧道、五线并行特大桥",
        '"五位一体"管控体系：安全零事故、质量优良、成本受控'
    ],
    "experiences": [
        {"period": "2008年7月 — 2021年1月", "company": "中铁隧道集团二处有限公司", "role": "技术员 → 工程部长 → 副总工程师", "summary": "先后参与多个隧道及地铁项目技术管理工作，从一线技术岗位逐步成长，积累了矿山法、盾构法、明挖法等隧道施工全工法技术经验，其间取得注册一级建造师执业资格。"},
        {"period": "2021年2月 — 2024年2月", "company": "南京地铁七号线D7-TA02标土建七工区清凉山站项目经理部", "role": "副经理（主持工作），负责项目全面工作", "summary": "主持南京地铁七号线清凉山站建设，项目合同额4.27亿元。通过优化围护结构、调整施工顺序、优化暗挖段支护参数等系列技术管理措施，实现责任成本盈利453.81万元。"},
        {"period": "2024年2月 — 2025年3月", "company": "北京轨道交通1号线支线工程土建施工09合同段项目经理部", "role": "项目经理", "summary": "主持北京地铁1号线支线体育场南街站建设，合同额2.44亿元。地下两层岛式车站，围护结构采用800mm地下连续墙+三道钢支撑。实现责任成本盈利115万元。"},
        {"period": "2025年3月 — 至今", "company": "新建深圳西丽站及相关工程XLSG-1标项目经理部", "role": "党支部书记、项目经理（兼任一分部工作）", "summary": "主持深圳市重点工程——新建深圳西丽站及相关工程XLSG-1标，合同总额32.76亿元（二处分劈19.72亿元），管理类别A类。项目涵盖西丽隧道（盾构法+矿山法）、塘朗山隧道、大沙河特大桥等重难点工程。"}
    ],
    "projects": [
        {"name": "南京地铁七号线清凉山站", "method": "矿山法", "role": "副经理（主持工作）", "period": "2021.02 — 2024.02", "metrics": [{"value": "4.27亿", "label": "合同额"}, {"value": "596.65m", "label": "车站总长"}, {"value": "303㎡", "label": "最大断面"}, {"value": "2300万", "label": "节约成本"}], "highlights": ["优化围护结构为钻孔灌注咬合桩，节约成本2300万元", "暗挖段优化拱盖法支护参数，节约工期约90天", "针对303㎡大跨断面，利用临时中立柱及增设L型吊点确保二衬钢筋稳定"]},
        {"name": "北京地铁1号线支线体育场南街站", "method": "明挖法", "role": "项目经理", "period": "2024.02 — 2025.03", "metrics": [{"value": "2.44亿", "label": "合同额"}, {"value": "229.2m", "label": "车站长度"}, {"value": "800mm", "label": "地下连续墙"}, {"value": "6%", "label": "成本盈利率"}], "highlights": ["双层三跨矩形框架结构，800mm地下连续墙+三道钢支撑围护体系", "实现工程成本盈利率6%，责任成本盈利率2%", "3个出入口+2个安全出口+2组风亭，功能布局完整"]},
        {"name": "新建深圳西丽站及相关工程XLSG-1标", "method": "盾构", "role": "党支部书记、项目经理", "period": "2025.03 — 至今", "metrics": [{"value": "32.76亿", "label": "合同总额"}, {"value": "19.72亿", "label": "二处分劈"}, {"value": "2025-2028", "label": "合同工期"}, {"value": "A类", "label": "管理类别"}], "highlights": ["深圳市重点工程，西丽隧道盾构法及矿山法下穿多处建构筑物", "塘朗山隧道小净距施工+燕尾段大断面，工法转换复杂", "大沙河特大桥5线7股并行、多处上跨市政道路"]}
    ],
    "education": [
        {"school": "兰州交通大学", "major": "土木工程", "degree": "全日制本科", "period": "2004.09 — 2008.07"},
        {"school": "清华大学", "major": "卓越领袖研修班", "degree": "研修深造", "period": "2025 — 2026.09（在读）"}
    ],
    "certifications": [
        {"name": "一级建造师（市政公用工程）", "year": "2019"},
        {"name": "职业项目经理（三级）", "year": "2025"}
    ],
    "closing": "以履约立身，以经营致远，以创新破局。",
    "skills": [
        {"name": "项目管理", "color": "navy", "pct": 95},
        {"name": "市场经营", "color": "gold", "pct": 85},
        {"name": "技术创新", "color": "blue", "pct": 80},
        {"name": "AI 数智化", "color": "red", "pct": 90}
    ],
    "ai_scenarios": [
        {"icon": "1", "name": "日报/周报自动生成", "desc": "3-5小时 → 15分钟"},
        {"icon": "2", "name": "会议纪要待办追踪", "desc": "1-2小时 → 15分钟"},
        {"icon": "3", "name": "招标情报每日监测", "desc": "2小时 → 12分钟"},
        {"icon": "4", "name": "施组方案智能起草", "desc": "编制周期缩短30%+"},
        {"icon": "5", "name": "工法专利报告起草", "desc": "周期缩短50%+"}
    ],
    "awards": [],
    "patents": [],
}


def extract_docx(filepath):
    """解析 .docx 文件"""
    from docx import Document
    doc = Document(filepath)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    tables = []
    for t in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        tables.append(rows)
    return text, tables


def extract_doc(filepath):
    """解析 .doc 文件（antiword + olefile 双重提取）"""
    text = ""
    # 方法1：antiword
    try:
        result = subprocess.run(["antiword", filepath], capture_output=True, text=True, timeout=15)
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout
    except Exception:
        pass

    # 方法2：olefile 直接读取 WordDocument 流（兜底）
    if not text:
        try:
            import olefile
            ole = olefile.OleFileIO(filepath)
            if ole.exists('WordDocument'):
                raw = ole.openstream('WordDocument').read()
                # 提取可打印字符序列（UTF-16LE 编码的中文文本）
                try:
                    decoded = raw.decode('utf-16-le', errors='ignore')
                    # 提取中文字符和常见标点
                    import re as re_mod
                    chunks = re_mod.findall(r'[一-鿿\d\w\s.,;:;，。；：、\n]{2,}', decoded)
                    text = '\n'.join(chunks)
                except:
                    pass
            ole.close()
        except Exception:
            pass

    if not text:
        text = "[提示] 无法提取文本，请上传 .docx 格式文件。"

    photos = []
    try:
        with open(filepath, 'rb') as f:
            raw = f.read()
        pos = 0
        while True:
            idx = raw.find(b'\xff\xd8\xff', pos)
            if idx == -1:
                break
            end = raw.find(b'\xff\xd9', idx)
            if end != -1:
                photos.append(('jpg', raw[idx:end + 2]))
                pos = end + 2
            else:
                pos = idx + 1
        pos = 0
        while True:
            idx = raw.find(b'\x89PNG\r\n\x1a\n', pos)
            if idx == -1:
                break
            end = raw.find(b'IEND\xaeB`\x82', idx)
            if end != -1:
                photos.append(('png', raw[idx:end + 8]))
                pos = end + 8
            else:
                pos = idx + 1
    except Exception:
        pass

    return text, photos


def auto_extract_data(text, tables, photos):
    """从上传文件中提取基本个人信息。仅提取确定字段，不生成任何合成文本。"""
    import copy
    data = copy.deepcopy(DEFAULT_DATA)

    # 合并表格内容到文本
    for table in tables:
        for row in table:
            for cell in row:
                if cell and cell not in text:
                    text += "\n" + cell

    # 个人信息（兼容有冒号和无冒号两种格式，antiword 输出可能没有冒号）
    simple_patterns = {
        "name": [
            r'姓\s*名\s*[：:：\s]+([一-鿿]{2,4})',
            r'姓\s*名\s*[：:：]?\s*([一-鿿]{2,4})',
        ],
        "gender": [
            r'性\s*别\s*[：:：\s]+([一-鿿男女]{1,2})',
            r'性\s*别\s*[：:：]?\s*([一-鿿男女]{1,2})',
        ],
        "birth": [
            r'出生[年月日日期]*\s*[：:：\s]+(\d{4}\s*[年.\-/]\s*\d{1,2})',
            r'出生[年月日日期]*\s*[：:：]?\s*(\d{4}\s*[年.\-/]\s*\d{1,2})',
        ],
        "political": [
            r'政治[面貌]*\s*[：:：\s]+([一-鿿]{2,8})',
            r'政治[面貌]*\s*[：:：]?\s*([一-鿿]{2,8})',
        ],
        "hometown": [
            r'籍\s*贯\s*[：:：\s]+([一-鿿]{2,10})',
            r'籍\s*贯\s*[：:：]?\s*([一-鿿]{2,10})',
        ],
        "phone": [
            r'(?:手机|电话|联系方式)\s*[：:：\s]+(\d{11})',
            r'(?:手机|电话|联系方式)\s*[：:：]?\s*(\d{11})',
        ],
        "location": [
            r'(?:工作地|所在地|现住址|地区)\s*[：:：\s]+(\S{3,20})',
            r'(?:工作地|所在地|现住址|地区)\s*[：:：]?\s*(\S{3,20})',
        ],
    }
    for field, patterns in simple_patterns.items():
        for pat in patterns:
            m = re.search(pat, text)
            if m:
                val = m.group(1).strip()
                if val:
                    data[field] = val
                    break

    # 照片
    if photos:
        fmt, img_data = photos[0]
        b64 = base64.b64encode(img_data).decode()
        data["photo"] = f"data:image/{fmt};base64,{b64}"

    return data


def render_html(data):
    """Jinja2 渲染"""
    from jinja2 import Environment
    env = Environment()
    template = env.from_string(RESUME_TEMPLATE)
    return template.render(**data)


def html_to_pdf(html_path):
    """Chrome headless → PDF（本地可用时）"""
    pdf_path = html_path.replace('.html', '.pdf')
    chrome_paths = [
        "google-chrome", "chromium-browser", "chromium",
        "/usr/bin/chromium", "/usr/bin/chromium-browser",
        "C:/Program Files/Google/Chrome/Application/chrome.exe",
        "C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe",
    ]
    for chrome in chrome_paths:
        # 检查命令是否存在
        if os.path.exists(chrome) or shutil_which(chrome):
            try:
                file_url = f"file:///{html_path.replace(chr(92), '/')}"
                subprocess.run([chrome, "--headless", "--disable-gpu",
                    f"--print-to-pdf={pdf_path}", "--no-pdf-header-footer", file_url],
                    check=True, timeout=30, capture_output=True)
                return pdf_path
            except Exception:
                continue
    return None


def shutil_which(cmd):
    """跨平台 which"""
    import shutil
    return shutil.which(cmd)


# ── UI ──
st.title("📄 央企项目经理竞聘简历生成器")
st.caption("上传 .doc 或 .docx 文件 → 自动提取 → 编辑确认 → 一键生成 PDF")

tab1, tab2, tab3 = st.tabs(["📤 上传与提取", "✏️ 编辑数据", "📄 生成下载"])

# ── Tab 1: 上传与提取 ──
with tab1:
    st.subheader("第一步：上传文件")
    st.caption("上传 .doc 或 .docx 认定表 → 自动提取照片和基本信息 → 再到「编辑数据」手动完善")
    uploaded = st.file_uploader("选择 .doc 或 .docx 文件", type=["doc", "docx"], key="upload")

    # 示例模板按钮
    col_a, col_b = st.columns([1, 1])
    with col_a:
        if st.button("📋 使用示例模板（项目经理竞聘范本）", use_container_width=True):
            import copy
            st.session_state['data'] = copy.deepcopy(SAMPLE_DATA)
            st.success("✅ 示例模板已加载！切换到「编辑数据」查看和修改。")
    with col_b:
        if st.button("🗑️ 清空所有数据", use_container_width=True):
            import copy
            st.session_state['data'] = copy.deepcopy(DEFAULT_DATA)
            st.success("已清空，可重新上传或使用模板。")

    if uploaded:
        # 仅首次上传或换文件时重新提取，避免覆盖用户手动编辑
        file_key = f"{uploaded.name}_{uploaded.size}"
        if st.session_state.get('_last_extracted_file') != file_key:
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded.name.split('.')[-1]}") as tmp:
                tmp.write(uploaded.read())
                tmp_path = tmp.name

            if uploaded.name.endswith('.docx'):
                text, tables = extract_docx(tmp_path)
                photos = []
            else:
                text, photos = extract_doc(tmp_path)
                tables = []

            os.unlink(tmp_path)
            st.session_state['extracted_text'] = text
            st.session_state['extracted_photos'] = photos

            data = auto_extract_data(text, tables, photos)
            st.session_state['data'] = data
            st.session_state['_last_extracted_file'] = file_key

        # 显示结果（从 session_state 读取，避免重复提取）
        data = st.session_state.get('data', {})
        text = st.session_state.get('extracted_text', '')
        photos = st.session_state.get('extracted_photos', [])

        if text:
            st.text_area("📝 提取的原始文本（供核对参考）", text[:3000], height=250)
        if photos:
            for i, (fmt, img_data) in enumerate(photos):
                b64 = base64.b64encode(img_data).decode()
                st.image(f"data:image/{fmt};base64,{b64}", caption=f"照片 {i+1}", width=120)

        found = []
        if data.get('name'): found.append(f"姓名={data['name']}")
        if data.get('gender'): found.append(f"性别={data['gender']}")
        if data.get('birth'): found.append(f"出生={data['birth']}")
        if data.get('political'): found.append(f"政治面貌={data['political']}")
        if data.get('phone'): found.append(f"手机={data['phone']}")
        if data.get('location'): found.append(f"工作地={data['location']}")
        if data.get('photo'): found.append("照片=已提取")
        if found:
            st.success(f"✅ 已自动提取并载入表单：{' | '.join(found)}")
            st.info("👉 请切换到上方「✏️ 编辑数据」标签页，查看并补全工作经历、项目经验等内容。")
        else:
            st.warning("⚠️ 未能自动提取信息。建议：① 切换到「编辑数据」手动填写 ② 或点击「使用示例模板」参考格式。")

# ── Tab 2: 编辑数据 ──
with tab2:
    if 'data' not in st.session_state:
        # 首次打开：空白表单（不加载任何模板，避免数据污染）
        import copy
        st.session_state['data'] = copy.deepcopy(DEFAULT_DATA)

    data = st.session_state['data']

    # 提取摘要
    exp_n = len(data.get('experiences', []))
    proj_n = len(data.get('projects', []))
    edu_n = len(data.get('education', []))
    cert_n = len(data.get('certifications', []))
    has_photo = bool(data.get('photo', ''))
    fields_filled = sum(1 for k, v in data.items() if v and k not in ('track_business', 'track_delivery', 'ai_scenarios', 'closing', 'skills', 'experiences', 'projects', 'education', 'certifications', 'awards', 'patents', 'stats', 'photo'))
    stats_filled = sum(1 for v in data.get('stats', {}).values() if v)
    if fields_filled > 0 or exp_n > 0:
        cols = st.columns(6)
        cols[0].metric("个人信息", fields_filled)
        cols[1].metric("数据仪表盘", stats_filled)
        cols[2].metric("工作经历", exp_n)
        cols[3].metric("项目经验", proj_n)
        cols[4].metric("教育背景", edu_n)
        cols[5].metric("证书", cert_n)
        if has_photo:
            st.caption("📷 已提取照片")
        st.divider()

    col1, col2, col3 = st.columns(3)
    with col1:
        data['name'] = st.text_input("姓名", data.get('name', ''))
        data['title'] = st.text_input("当前职务", data.get('title', ''))
        data['gender'] = st.text_input("性别", data.get('gender', ''))
        data['birth'] = st.text_input("出生年月", data.get('birth', ''))
        data['location'] = st.text_input("工作地", data.get('location', ''))
    with col2:
        data['phone'] = st.text_input("手机", data.get('phone', ''))
        data['email'] = st.text_input("邮箱", data.get('email', ''))
        data['hometown'] = st.text_input("籍贯", data.get('hometown', ''))
        data['political'] = st.text_input("政治面貌", data.get('political', ''))
    with col3:
        data['stats']['experience_years'] = st.text_input("从业年限", data['stats'].get('experience_years', ''))
        data['stats']['total_contract'] = st.text_input("管理合同额(亿)", data['stats'].get('total_contract', ''))
        data['stats']['project_count'] = st.text_input("项目数量", data['stats'].get('project_count', ''))
        data['stats']['team_size'] = st.text_input("团队规模", data['stats'].get('team_size', ''))

    data['one_liner'] = st.text_area("一句话定位", data.get('one_liner', ''), height=60)

    st.divider()
    st.subheader("工作经历")
    exp_json = st.text_area("JSON 格式", json.dumps(data.get('experiences', []), ensure_ascii=False, indent=2), height=200)
    try:
        data['experiences'] = json.loads(exp_json)
    except Exception:
        pass

    st.divider()
    st.subheader("项目经验")
    proj_json = st.text_area("JSON 格式", json.dumps(data.get('projects', []), ensure_ascii=False, indent=2), height=200, key="proj")
    try:
        data['projects'] = json.loads(proj_json)
    except Exception:
        pass

    st.divider()
    col4, col5 = st.columns(2)
    with col4:
        st.subheader("教育背景")
        edu_json = st.text_area("JSON", json.dumps(data.get('education', []), ensure_ascii=False, indent=2), height=120, key="edu")
        try:
            data['education'] = json.loads(edu_json)
        except Exception:
            pass
    with col5:
        st.subheader("资质证书")
        cert_json = st.text_area("JSON", json.dumps(data.get('certifications', []), ensure_ascii=False, indent=2), height=120, key="cert")
        try:
            data['certifications'] = json.loads(cert_json)
        except Exception:
            pass

    photo_val = data.get('photo', '')
    data['photo'] = st.text_input("照片 (base64 data URI 或留空)", photo_val[:80] + "..." if len(photo_val) > 80 else photo_val)

# ── Tab 3: 生成下载 ──
with tab3:
    st.subheader("生成简历")

    if st.button("🔨 生成 HTML", use_container_width=True, type="primary"):
        with st.spinner("渲染中..."):
            html = render_html(st.session_state['data'])
            out_dir = Path(tempfile.mkdtemp())
            html_path = out_dir / "resume.html"
            html_path.write_text(html, encoding='utf-8')
            st.session_state['html'] = html
            st.session_state['html_path'] = str(html_path)
            st.success("HTML 生成成功！")

    if 'html' in st.session_state:
        with st.expander("📄 HTML 预览", expanded=False):
            st.components.v1.html(st.session_state['html'], height=800, scrolling=True)

        st.download_button("⬇️ 下载 HTML", st.session_state['html'], "个人简历.html", "text/html", use_container_width=True)

        if st.button("🖨️ 转换为 PDF", use_container_width=True):
            with st.spinner("Chrome headless 转换中（云端可能不可用）..."):
                pdf_path = html_to_pdf(st.session_state['html_path'])
                if pdf_path and os.path.exists(pdf_path):
                    with open(pdf_path, 'rb') as f:
                        st.session_state['pdf_data'] = f.read()
                    st.success("PDF 生成成功！")
                else:
                    st.warning("⚠️ 服务器端 Chrome 不可用，请下载 HTML 后用浏览器打开 → Ctrl+P → 另存为 PDF。")

        if 'pdf_data' in st.session_state:
            st.download_button("⬇️ 下载 PDF", st.session_state['pdf_data'], "个人简历.pdf", "application/pdf", use_container_width=True, type="primary")

    st.divider()
    st.caption("💡 提示：PDF 转换需要 Chrome/Chromium。云端环境如转换失败，请下载 HTML 后浏览器打印为 PDF。")

# ── Footer ──
st.divider()
st.caption("V4 · IBM×Apple×Stripe 设计语言 · 央企项目经理竞聘专用 · Powered by Streamlit Cloud")
